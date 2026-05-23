"""Генератор docx-отчётов по лабораторным курса «Криптографические протоколы».

Шаблон титульного листа повторяет эталон из методов и средств защиты информации.
Преподаватель — Дубровина А.С., год — 2026, студент — Ковалев Д.П. ВКБ43.

Формулы вставляются как OMML-объекты Word (математическая вставка), а не как
тексты — функция `add_math` принимает OMML-XML и инжектит его прямо в параграф.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree

LOGO_PATH = Path(__file__).resolve().parents[1] / "docs" / "assets" / "dstu_logo.png"

STUDENT_NAME = "Ковалев Данил Петрович"
STUDENT_GROUP = "ВКБ43"
TEACHER = "Дубровина А.С."
DISCIPLINE = "Криптографические протоколы"
CITY = "Ростов-на-Дону"
YEAR = "2026"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(14)
MONO_FONT = "Courier New"
MONO_SIZE = Pt(10)

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NSMAP = {"m": M_NS}


@dataclass
class LabMeta:
    number: int
    title: str  # тема работы
    variant: int | None = None


def _setup_default_style(doc: _Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)


def _set_margins(doc: _Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def _add_centered(
    doc: _Document, text: str, *, bold: bool = False, size: Pt | None = None
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = size
    run.font.name = BODY_FONT


def _add_right(doc: _Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT


def _add_logo(doc: _Document) -> None:
    """Логотип ДГТУ по центру в начале титульника."""
    if not LOGO_PATH.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(0.8))


def add_title_page(doc: _Document, meta: LabMeta) -> None:
    """Добавляет стандартный титульник ДГТУ. Помещается на одну страницу A4."""
    _add_logo(doc)
    _add_centered(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", bold=True)
    _add_centered(doc, "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ", bold=True)
    _add_centered(doc, "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ", bold=True)
    _add_centered(doc, "«ДОНСКОЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»", bold=True)
    _add_centered(doc, "(ДГТУ)", bold=True)
    doc.add_paragraph()
    _add_centered(doc, "Факультет «Информатика и вычислительная техника»")
    _add_centered(doc, "Кафедра «Кибербезопасность информационных систем»")
    for _ in range(3):
        doc.add_paragraph()
    _add_centered(doc, f"Лабораторная работа №{meta.number}", bold=True, size=Pt(16))
    _add_centered(doc, f"по дисциплине: «{DISCIPLINE}»")
    variant_suffix = (
        f" (вариант №{meta.variant})" if meta.variant is not None else ""
    )
    _add_centered(doc, f"На тему «{meta.title}»{variant_suffix}", bold=True)
    for _ in range(3):
        doc.add_paragraph()
    _add_right(doc, f"Выполнил обучающийся гр. {STUDENT_GROUP}")
    _add_right(doc, STUDENT_NAME)
    doc.add_paragraph()
    _add_right(doc, "Проверила:")
    _add_right(doc, TEACHER)
    # Город и год — внизу титульника. При интервале 1.5 семь пустых параграфов уже
    # выталкивают строки на следующую страницу, поэтому ограничиваемся одним и
    # полагаемся на page break в самом скрипте лабы, который ставится после
    # add_title_page().
    doc.add_paragraph()
    _add_centered(doc, CITY)
    _add_centered(doc, YEAR)


def add_heading(doc: _Document, text: str, *, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(16 if level == 1 else 14)


def add_para(doc: _Document, text: str, *, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE


def add_label(doc: _Document, text: str) -> None:
    """Подпись к рисунку/листингу — по центру."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = BODY_FONT
    run.font.size = Pt(13)


def add_listing(doc: _Document, code: str, caption: str | None = None) -> None:
    """Листинг кода/вывода — моноширинный, single line spacing (исключение из правила 1.5)."""
    for line in code.rstrip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line if line else " ")
        run.font.name = MONO_FONT
        run.font.size = MONO_SIZE
        run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
    if caption:
        add_label(doc, caption)


def add_bullets(doc: _Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE


def add_page_break(doc: _Document) -> None:
    """Удобная обёртка для разрыва страницы."""
    doc.add_page_break()


def add_qa(doc: _Document, n: int, question: str, answer: str) -> None:
    """Контрольный вопрос + ответ. Вопрос жирный, ответ — обычный."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run_q = p.add_run(f"Вопрос {n}. {question}")
    run_q.bold = True
    run_q.font.name = BODY_FONT
    run_q.font.size = BODY_SIZE

    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_a = pa.add_run(f"Ответ. {answer}")
    run_a.font.name = BODY_FONT
    run_a.font.size = BODY_SIZE


def add_math(doc: _Document, omml_xml: str, *, centered: bool = True) -> None:
    """Вставить математическую формулу как OMML-объект.

    `omml_xml` должен быть полным <m:oMath>...</m:oMath> или <m:oMathPara>...</m:oMathPara>.
    Word отрисует это как формулу, не как простой текст.
    """
    p = doc.add_paragraph()
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    elem = etree.fromstring(omml_xml)
    p._p.append(elem)


def _omml_r(text: str, *, italic: bool = False) -> str:
    """OMML <m:r> с текстом."""
    style = '<m:rPr><m:sty m:val="i"/></m:rPr>' if italic else ""
    # Текст экранируем минимально (для нашего случая опасных символов нет).
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return f'<m:r xmlns:m="{M_NS}">{style}<m:t xml:space="preserve">{safe}</m:t></m:r>'


def omml_inline(parts: list[str]) -> str:
    """Собрать inline-математическую вставку из готовых OMML-кусков (одна формула)."""
    body = "".join(parts)
    return f'<m:oMath xmlns:m="{M_NS}">{body}</m:oMath>'


def omml_display(parts: list[str]) -> str:
    """Display-форма (на отдельной строке)."""
    body = "".join(parts)
    return (
        f'<m:oMathPara xmlns:m="{M_NS}">'
        f"<m:oMathParaPr><m:jc m:val=\"center\"/></m:oMathParaPr>"
        f"<m:oMath>{body}</m:oMath></m:oMathPara>"
    )


# ----- удобные конструкторы для часто встречающихся формул -----


def m_text(text: str, *, italic: bool = True) -> str:
    return _omml_r(text, italic=italic)


def m_op(symbol: str) -> str:
    """Бинарный/унарный оператор: +, −, =, ≡, ·, и т.п. — не курсивный."""
    return _omml_r(symbol, italic=False)


def m_sup(base: str, exp: str) -> str:
    return (
        f'<m:sSup xmlns:m="{M_NS}">'
        f"<m:e>{base}</m:e>"
        f"<m:sup>{exp}</m:sup>"
        f"</m:sSup>"
    )


def m_sub(base: str, sub: str) -> str:
    return (
        f'<m:sSub xmlns:m="{M_NS}">'
        f"<m:e>{base}</m:e>"
        f"<m:sub>{sub}</m:sub>"
        f"</m:sSub>"
    )


def m_frac(num: str, den: str) -> str:
    return (
        f'<m:f xmlns:m="{M_NS}">'
        f"<m:num>{num}</m:num>"
        f"<m:den>{den}</m:den>"
        f"</m:f>"
    )


def m_sum(lower: str, upper: str, body: str) -> str:
    """∑_{lower}^{upper} body — большой сумматор."""
    return (
        f'<m:nary xmlns:m="{M_NS}">'
        f'<m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/></m:naryPr>'
        f"<m:sub>{lower}</m:sub>"
        f"<m:sup>{upper}</m:sup>"
        f"<m:e>{body}</m:e>"
        f"</m:nary>"
    )


def m_prod(lower: str, upper: str, body: str) -> str:
    """∏_{lower}^{upper} body."""
    return (
        f'<m:nary xmlns:m="{M_NS}">'
        f'<m:naryPr><m:chr m:val="∏"/><m:limLoc m:val="undOvr"/></m:naryPr>'
        f"<m:sub>{lower}</m:sub>"
        f"<m:sup>{upper}</m:sup>"
        f"<m:e>{body}</m:e>"
        f"</m:nary>"
    )


def make_doc() -> _Document:
    doc = Document()
    _setup_default_style(doc)
    _set_margins(doc)
    return doc


def save(doc: _Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
