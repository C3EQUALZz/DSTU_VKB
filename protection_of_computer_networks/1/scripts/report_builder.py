"""Генератор docx-отчётов по курсу «Защита компьютерных сетей».

Шаблон титульника повторяет образец курса «Криптографические протоколы» —
ДГТУ, кафедра КБИС, та же типографика (Times New Roman 14, поля 3/1.5/2/2,
полуторный интервал, отступ 1.25 см). Преподаватель — Болдырихин Н.В.

Эту копию вынесли отдельно от cryptographic_protocols/scripts/report_builder.py,
чтобы дисциплина и ФИО преподавателя жили рядом с самими отчётами курса.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt, RGBColor

LOGO_PATH = Path(__file__).resolve().parents[1] / "docs" / "assets" / "dstu_logo.png"

STUDENT_NAME = "Ковалев Данил Петрович"
STUDENT_GROUP = "ВКБ43"
TEACHER = "Болдырихин Н.В."
DISCIPLINE = "Защита компьютерных сетей"
CITY = "Ростов-на-Дону"
YEAR = "2026"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(14)
MONO_FONT = "Courier New"
MONO_SIZE = Pt(10)


@dataclass
class LabMeta:
    number: int
    title: str
    variant: int | None = None


def _setup_default_style(doc: _Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)


def _set_first_page_footer_with_city_year(doc: _Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.first_page_footer
    paragraphs = list(footer.paragraphs) or [footer.add_paragraph()]
    first = paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.first_line_indent = Cm(0)
    first.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    first.clear()
    run1 = first.add_run(CITY)
    run1.font.name = BODY_FONT
    run1.font.size = BODY_SIZE

    second = footer.add_paragraph()
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second.paragraph_format.first_line_indent = Cm(0)
    second.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run2 = second.add_run(YEAR)
    run2.font.name = BODY_FONT
    run2.font.size = BODY_SIZE


def _set_margins(doc: _Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def _add_centered(
    doc: _Document, text: str, *, bold: bool = False, size: Pt | None = None
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = size
    run.font.name = BODY_FONT


def _add_right(doc: _Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT


def _add_blank_single(doc: _Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _add_logo(doc: _Document) -> None:
    if not LOGO_PATH.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(0.8))


def add_title_page(doc: _Document, meta: LabMeta) -> None:
    _add_logo(doc)
    _add_centered(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", bold=True)
    _add_centered(doc, "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ", bold=True)
    _add_centered(doc, "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ", bold=True)
    _add_centered(doc, "«ДОНСКОЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»", bold=True)
    _add_centered(doc, "(ДГТУ)", bold=True)
    _add_blank_single(doc)
    _add_centered(doc, "Факультет «Информатика и вычислительная техника»")
    _add_centered(doc, "Кафедра «Кибербезопасность информационных систем»")
    for _ in range(4):
        _add_blank_single(doc)
    _add_centered(doc, f"Лабораторная работа №{meta.number}", bold=True, size=Pt(16))
    _add_centered(doc, f"по дисциплине: «{DISCIPLINE}»")
    variant_suffix = (
        f" (вариант №{meta.variant})" if meta.variant is not None else ""
    )
    _add_centered(doc, f"На тему «{meta.title}»{variant_suffix}", bold=True)
    for _ in range(5):
        _add_blank_single(doc)
    _add_right(doc, f"Выполнил обучающийся гр. {STUDENT_GROUP}")
    _add_right(doc, STUDENT_NAME)
    _add_blank_single(doc)
    _add_right(doc, "Проверил:")
    _add_right(doc, TEACHER)
    _set_first_page_footer_with_city_year(doc)


def add_heading(doc: _Document, text: str, *, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(16 if level == 1 else 14)


def add_para(doc: _Document, text: str, *, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE


def add_label(doc: _Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = BODY_FONT
    run.font.size = Pt(13)


def add_listing(doc: _Document, code: str, caption: str | None = None) -> None:
    """Листинг кода/вывода: моноширинный, single line spacing."""
    for line in code.rstrip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line if line else " ")
        run.font.name = MONO_FONT
        run.font.size = MONO_SIZE
        run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
    if caption:
        add_label(doc, caption)


def add_bullets(doc: _Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE


def add_page_break(doc: _Document) -> None:
    doc.add_page_break()


def add_qa(doc: _Document, n: int, question: str, answer: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run_q = p.add_run(f"Вопрос {n}. {question}")
    run_q.bold = True
    run_q.font.name = BODY_FONT
    run_q.font.size = BODY_SIZE

    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_a = pa.add_run(f"Ответ. {answer}")
    run_a.font.name = BODY_FONT
    run_a.font.size = BODY_SIZE


def make_doc() -> _Document:
    doc = Document()
    _setup_default_style(doc)
    _set_margins(doc)
    return doc


def save(doc: _Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
