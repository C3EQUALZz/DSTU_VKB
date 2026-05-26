"""Генератор docx-отчётов курса «Стеганография» (ДГТУ, ВКБ).

Шаблон титульного листа повторяет образцы 2025 г. по дисциплине
«Математические методы сокрытия и маскирования информации».
Преподаватель по умолчанию — доц. Сафарьян О.А., группа ВКБ51,
студент — Ковалев Данил Петрович (правится в STUDENT_NAME).
"""

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Inches, Pt, RGBColor

LOGO_PATH = (
    Path(__file__).resolve().parents[1] / "docs" / "assets" / "dstu_logo.png"
)

STUDENT_NAME = "Ковалев Данил Петрович"
STUDENT_GROUP = "ВКБ51"
TEACHER_DEFAULT = "доц. Сафарьян О.А."
DISCIPLINE = "Математические методы сокрытия и маскирования информации"
CITY = "Ростов-на-Дону"
YEAR = "2025"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(14)
MONO_FONT = "Courier New"
MONO_SIZE = Pt(10)


@dataclass
class LabMeta:
    number: int
    title: str
    work_kind: str = "Практическая работа"
    variant: int | None = None
    teacher: str = TEACHER_DEFAULT


def _setup_default_style(doc: _Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)


def _set_first_page_footer_with_city_year(doc: _Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.first_page_footer
    paragraphs = list(footer.paragraphs) or [footer.add_paragraph()]
    first = paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.first_line_indent = Cm(0)
    first.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    first.clear()
    run1 = first.add_run(CITY)
    run1.font.name = BODY_FONT
    run1.font.size = BODY_SIZE

    second = footer.add_paragraph()
    second.alignment = WD_ALIGN_PARAGRAPH.CENTER
    second.paragraph_format.first_line_indent = Cm(0)
    second.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run2 = second.add_run(f"{YEAR} г.")
    run2.font.name = BODY_FONT
    run2.font.size = BODY_SIZE


def _set_margins(doc: _Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def _add_centered(
    doc: _Document,
    text: str,
    *,
    bold: bool = False,
    size: Pt | None = None,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = size
    run.font.name = BODY_FONT


def _add_right(doc: _Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.font.name = BODY_FONT


def _add_blank_single(doc: _Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _add_logo(doc: _Document) -> None:
    if not LOGO_PATH.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(0.8))


def add_title_page(doc: _Document, meta: LabMeta) -> None:
    _add_logo(doc)
    _add_centered(
        doc,
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
        bold=True,
    )
    _add_centered(
        doc, "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ", bold=True,
    )
    _add_centered(
        doc, "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ", bold=True,
    )
    _add_centered(
        doc, "«ДОНСКОЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ»", bold=True,
    )
    _add_centered(doc, "(ДГТУ)", bold=True)
    _add_blank_single(doc)
    _add_centered(doc, "Факультет «Информатика и вычислительная техника»")
    _add_centered(doc, "Кафедра «Кибербезопасность информационных систем»")
    for _ in range(4):
        _add_blank_single(doc)
    _add_centered(doc, "ОТЧЁТ", bold=True, size=Pt(16))
    _add_centered(
        doc,
        f"{meta.work_kind} № {meta.number}",
        bold=True,
        size=Pt(16),
    )
    _add_centered(doc, f"по дисциплине «{DISCIPLINE}»")
    variant_suffix = (
        f" (вариант №{meta.variant})" if meta.variant is not None else ""
    )
    _add_centered(doc, f"На тему «{meta.title}»{variant_suffix}", bold=True)
    for _ in range(5):
        _add_blank_single(doc)
    _add_right(doc, f"Выполнил: студент группы {STUDENT_GROUP}")
    _add_right(doc, STUDENT_NAME)
    _add_blank_single(doc)
    _add_right(doc, "Проверил:")
    _add_right(doc, meta.teacher)
    _set_first_page_footer_with_city_year(doc)


def add_heading(doc: _Document, text: str, *, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(16 if level == 1 else 14)


def add_para(doc: _Document, text: str, *, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE


def add_label(doc: _Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = BODY_FONT
    run.font.size = Pt(13)


def add_listing(
    doc: _Document, code: str, caption: str | None = None,
) -> None:
    for line in code.rstrip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line if line else " ")
        run.font.name = MONO_FONT
        run.font.size = MONO_SIZE
        run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
    if caption:
        add_label(doc, caption)


def add_image(
    doc: _Document,
    image_path: Path | str,
    caption: str | None = None,
    *,
    width_cm: float = 15.5,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        add_label(doc, caption)


def add_bullets(doc: _Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE


def add_table_simple(
    doc: _Document,
    rows: list[list[str]],
    *,
    header: bool = True,
    caption: str | None = None,
) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(str(value))
            run.font.name = BODY_FONT
            run.font.size = Pt(12)
            if header and r_idx == 0:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        add_label(doc, caption)


def add_page_break(doc: _Document) -> None:
    doc.add_page_break()


def make_doc() -> _Document:
    doc = Document()
    _setup_default_style(doc)
    _set_margins(doc)
    return doc


def save(doc: _Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
