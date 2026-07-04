"""Запись контейнера-результата в docx по плану форматирования.

Строки плана (:attr:`FormattingPlan.line_lengths`) воссоздаются как
отдельные абзацы docx, чтобы визуально документ повторял исходный
контейнер (стих с разбивкой на строки), а не склеивался в один абзац.
Для каждого символа создаётся отдельный run: в его свойства (``w:rPr``)
добавляется элемент параметра сокрытия с предписанным значением и базовый
шрифт контейнера. Элементы ``w:rPr`` добавляются в порядке, заданном
схемой OOXML (``rFonts`` → ``color`` → ``spacing`` → ``w`` → ``sz`` →
``highlight``), поэтому документ корректно открывается в Word и читается
детектором ПР «декод» побитово согласованно.
"""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run

from steganography.domain.common.value_objects.formatting_param import (
    FormattingParam,
)
from steganography.domain.text_format_encode.ports.container_writer import (
    ContainerWriter,
)
from steganography.domain.text_format_encode.value_objects.char_formatting import (
    CharFormatting,
)
from steganography.domain.text_format_encode.value_objects.formatting_plan import (
    FormattingPlan,
)


def _append_val(rpr: object, tag: str, value: str) -> None:
    element = OxmlElement(tag)
    element.set(qn("w:val"), value)
    rpr.append(element)  # type: ignore[attr-defined]


class DocxContainerWriterImpl(ContainerWriter):
    """Реализация порта записи через python-docx + прямые OOXML-элементы."""

    def write(self, plan: FormattingPlan, path: Path) -> None:
        document = Document()
        line_lengths = plan.line_lengths or (len(plan.chars),)

        offset = 0
        for length in line_lengths:
            paragraph = document.add_paragraph()
            for char_formatting in plan.chars[offset : offset + length]:
                run = paragraph.add_run(char_formatting.char)
                self._apply(run, char_formatting, plan)
            offset += length

        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(path))

    @staticmethod
    def _apply(
        run: Run, formatting: CharFormatting, plan: FormattingPlan,
    ) -> None:
        param = formatting.param
        value = formatting.value
        rpr = run._element.get_or_add_rPr()  # noqa: SLF001

        # Порядок дочерних элементов соответствует схеме CT_RPr.
        if plan.font_name:
            fonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                fonts.set(qn(attr), plan.font_name)
            rpr.append(fonts)
        if param is FormattingParam.COLOR:
            _append_val(rpr, "w:color", value)
        if param is FormattingParam.SPACING:
            _append_val(rpr, "w:spacing", value)
        if param is FormattingParam.SCALE:
            _append_val(rpr, "w:w", value)

        # Размер: носитель бита (param SIZE) или базовый шрифт контейнера.
        size_value = value if param is FormattingParam.SIZE else plan.font_size
        if size_value:
            _append_val(rpr, "w:sz", size_value)
            _append_val(rpr, "w:szCs", size_value)

        if param is FormattingParam.HIGHLIGHT:
            _append_val(rpr, "w:highlight", value)
