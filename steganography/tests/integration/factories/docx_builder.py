"""Фабрики реальных docx-файлов для интеграционных тестов."""

from pathlib import Path

from docx import Document
from docx.shared import Pt


def build_docx_with_size_hiding(
    path: Path,
    cover: str,
    bits: str,
    *,
    size0_pt: float = 14.0,
    size1_pt: float = 14.5,
) -> Path:
    """Создать docx, где сообщение скрыто через размер шрифта.

    Каждый символ ``cover`` получает размер ``size1_pt`` для бита «1» и
    ``size0_pt`` для бита «0». Биты добиваются нулями до длины ``cover``.
    """
    doc = Document()
    paragraph = doc.add_paragraph()
    padding = "0" * max(0, len(cover) - len(bits))
    full_bits = (bits + padding)[: len(cover)]
    for ch, bit in zip(cover, full_bits, strict=True):
        run = paragraph.add_run(ch)
        run.font.size = Pt(size1_pt if bit == "1" else size0_pt)
    doc.save(str(path))
    return path


def build_uniform_docx(path: Path, text: str, *, size_pt: float = 14.0) -> Path:
    """Создать docx без сокрытия — все символы одинакового размера."""
    doc = Document()
    paragraph = doc.add_paragraph()
    for ch in text:
        paragraph.add_run(ch).font.size = Pt(size_pt)
    doc.save(str(path))
    return path
